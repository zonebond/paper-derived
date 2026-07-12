"""多格式文档输出层.

输出策略: Markdown 优先。
1. DocumentTree → doc_tree_to_markdown() → .md 文件 (始终写入磁盘)
2. .md 文件 → Pandoc → .docx / .pdf (高质量转换)
3. Pandoc 不可用时降级到内置转换器 (质量受限)

支持的格式: json, md, docx, pdf
"""

from __future__ import annotations

import json
import shutil
import sys
from pathlib import Path


# ── 支持的输出格式 ──────────────────────────────────────────────

OUTPUT_FORMATS = {"json", "md", "docx", "pdf"}

# 扩展名 → 格式映射
_EXT_TO_FORMAT = {
    ".json": "json",
    ".md": "md",
    ".markdown": "md",
    ".docx": "docx",
    ".pdf": "pdf",
}


def infer_format(output_path: str | Path) -> str:
    """从输出路径的扩展名推断输出格式.

    Args:
        output_path: 输出文件路径.

    Returns:
        格式字符串 (json, md, docx, pdf).

    Raises:
        ValueError: 不支持的扩展名.
    """
    suffix = Path(output_path).suffix.lower()
    fmt = _EXT_TO_FORMAT.get(suffix)
    if fmt is None:
        raise ValueError(
            f"无法从扩展名推断输出格式: {suffix}。"
            f"支持的扩展名: {', '.join(sorted(_EXT_TO_FORMAT.keys()))}。"
            f"请使用 --format 显式指定。"
        )
    return fmt


# ── Markdown 渲染 ────────────────────────────────────────────────


def doc_tree_to_markdown(doc_tree) -> str:
    """将 DocumentTree 递归渲染为单个 Markdown 字符串.

    Args:
        doc_tree: DocumentTree 对象 (models/document.py).

    Returns:
        Markdown 格式的完整文档文本.
    """
    parts: list[str] = []

    # 文档标题
    if doc_tree.title:
        parts.append(f"# {doc_tree.title}\n")

    # 递归渲染所有 Section
    for section in doc_tree.sections:
        parts.append(_render_section(section))

    return "\n\n".join(parts)


def _render_section(section, heading_level: int = 2) -> str:
    """递归渲染单个 Section 为 Markdown.

    Args:
        section: Section 对象.
        heading_level: 当前 Section 的 heading 级别。
            顶层 Section 默认从 ## (H2) 开始，文档标题占 # (H1)。
            子 Section 自动递增一级。
            不使用 section.level 字段，因为实际层级由嵌套结构决定。

    Returns:
        Markdown 格式的 Section 文本.
    """
    lines: list[str] = []

    prefix = "#" * heading_level

    lines.append(f"{prefix} {section.title}")

    if section.content:
        lines.append("")
        lines.append(section.content)

    # 渲染子 Section — 层级递增
    for child in section.children:
        lines.append("")
        lines.append(_render_section(child, heading_level=min(heading_level + 1, 6)))

    # 如果有 lineage 信息, 添加来源标注
    if section.lineage and section.status == "generated":
        sources = [f"`{l.input_id}`" for l in section.lineage if l.input_id]
        if sources:
            lines.append("")
            lines.append(f"*[来源: {', '.join(sources)}]*")

    # 状态标记
    if section.status == "placeholder":
        lines.append("")
        lines.append(f"> ⚠️ 此章节为占位内容，需要补充输入资料。")
        if section.hints:
            for hint in section.hints:
                lines.append(f"> 💡 {hint}")

    return "\n".join(lines)


# ── Pandoc 转换 ─────────────────────────────────────────────────


def _ensure_pandoc() -> bool:
    """确保 pypandoc 能找到 Pandoc 二进制.

    搜索顺序:
    1. PyInstaller 打包模式: sys._MEIPASS 下的 pandoc
    2. PYPANDOC_PANDOC 环境变量
    3. pypandoc 自行搜索 (PATH, pypandoc-binary 等)

    Returns:
        True 找到 Pandoc, False 未找到.
    """
    import os
    import sys

    # PyInstaller onefile: pandoc 打包在 _MEIPASS 中
    meipass = getattr(sys, '_MEIPASS', None)
    if meipass:
        pandoc_in_bundle = os.path.join(meipass, 'pandoc')
        if os.path.isfile(pandoc_in_bundle):
            os.environ['PYPANDOC_PANDOC'] = pandoc_in_bundle
            return True

    # 已有环境变量或 pypandoc 能自行找到
    try:
        import pypandoc
        pypandoc.get_pandoc_path()
        return True
    except (ImportError, OSError):
        pass

    return False


def _convert_via_pandoc(md_path: Path, output_path: Path, to_format: str, title: str = "") -> bool:
    """用 Pandoc 从 Markdown 文件转换到目标格式.

    Args:
        md_path: 已写入磁盘的 .md 文件路径.
        output_path: 目标输出路径.
        to_format: 'docx' 或 'pdf'.
        title: 文档标题 (用于元数据).

    Returns:
        True 转换成功, False Pandoc 不可用.
    """
    if not _ensure_pandoc():
        return False

    try:
        import pypandoc
    except ImportError:
        return False

    try:
        extra_args: list[str] = []

        if to_format == "docx":
            # DOCX: Pandoc 原生支持，无需额外引擎
            pass

        elif to_format == "pdf":
            engine, engine_args = _pandoc_pdf_engine()
            if not engine:
                # 没有可用的 PDF 引擎，降级
                return False
            extra_args.append(f"--pdf-engine={engine}")
            extra_args.extend(engine_args)

        # 添加标题元数据
        if title:
            extra_args.extend(["-M", f"title={title}"])

        pypandoc.convert_file(
            str(md_path),
            to_format,
            format="md",
            outputfile=str(output_path),
            extra_args=extra_args if extra_args else None,
        )
        return True

    except OSError:
        # Pandoc 二进制不存在
        return False
    except Exception:
        # Pandoc 调用失败 (其他原因)
        return False


def _pandoc_pdf_engine() -> tuple[str, list[str]]:
    """检测可用的 Pandoc PDF 引擎.

    优先级: typst > weasyprint > xelatex

    Returns:
        (engine_name, extra_args) — 引擎不可用时返回 ('', []).
    """
    # 1. typst — 小体积 (~15MB), 原生中文支持, 无需 LaTeX
    if shutil.which("typst"):
        return "typst", ["-V", "lang=zh"]

    # 2. weasyprint — 纯 Python, pip install 即可
    try:
        import weasyprint  # noqa: F401
        return "weasyprint", []
    except ImportError:
        pass

    # 3. xelatex — 最重 (~4GB TeX Live), 但中文排版质量最高
    if shutil.which("xelatex"):
        import platform
        if platform.system() == "Darwin":
            return "xelatex", ["-V", "mainfont=PingFang SC"]
        return "xelatex", ["-V", "mainfont=Noto Sans CJK SC"]

    # 4. 全部不可用
    return "", []


# ── 主写入函数 ───────────────────────────────────────────────────


def write_document(doc_tree, output_path: str | Path, fmt: str | None = None) -> Path:
    """将 DocumentTree 写入指定格式文件.

    始终先写 Markdown 文件到磁盘，再从 Markdown 转换为目标格式。
    Markdown 文件与目标文件同目录、同名、.md 扩展名。

    Args:
        doc_tree: DocumentTree 对象.
        output_path: 输出文件路径.
        fmt: 输出格式 (json, md, docx, pdf)。若为 None, 从扩展名推断.

    Returns:
        实际写入的目标文件路径.

    Raises:
        ValueError: 不支持的格式.
    """
    path = Path(output_path).expanduser().resolve()
    path.parent.mkdir(parents=True, exist_ok=True)

    if fmt is None:
        fmt = infer_format(path)

    if fmt not in OUTPUT_FORMATS:
        raise ValueError(
            f"不支持的输出格式: {fmt}。支持的格式: {', '.join(sorted(OUTPUT_FORMATS))}"
        )

    # 1. 生成 Markdown 文本
    md_text = doc_tree_to_markdown(doc_tree)
    title = doc_tree.title if hasattr(doc_tree, "title") else ""

    # 2. JSON 格式直接输出 DocumentTree JSON（不需要 Markdown 中间产物）
    if fmt == "json":
        data = doc_tree.to_dict() if hasattr(doc_tree, "to_dict") else doc_tree
        path.write_text(
            json.dumps(data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        # JSON 模式也保留一份 Markdown
        md_path = path.with_suffix(".md")
        md_path.write_text(md_text, encoding="utf-8")
        return path

    # 3. 始终写 Markdown 文件
    md_path = path.with_suffix(".md")
    md_path.write_text(md_text, encoding="utf-8")

    # 4. Markdown 格式 — 已写入 md_path，返回实际写入的路径
    if fmt == "md":
        # 如果输出路径本身是 .md，md_path == path，直接返回
        # 如果输出路径是其他扩展名（如 .txt），内容在 md_path
        return md_path

    # 5. DOCX / PDF — 从 Markdown 文件转换
    success = _convert_via_pandoc(md_path, path, fmt, title=title)
    if success:
        return path

    # 6. Pandoc 不可用 → 降级到内置转换器
    print(
        "⚠️  Pandoc 不可用，使用内置转换器（输出质量受限）。\n"
        "   建议安装 Pandoc: https://pandoc.org/installing.html",
        file=sys.stderr,
    )
    if fmt == "docx":
        _markdown_to_docx_fallback(md_text, path, title)
    elif fmt == "pdf":
        _markdown_to_pdf_fallback(md_text, path, title)

    return path


def write_string(content: str, output_path: str | Path, fmt: str | None = None, title: str = "") -> Path:
    """将纯文本/Markdown 字符串写入指定格式文件.

    始终先写 Markdown 文件到磁盘，再从 Markdown 转换为目标格式。

    Args:
        content: 文本内容 (Markdown 格式).
        output_path: 输出文件路径.
        fmt: 输出格式 (md, docx, pdf)。若为 None, 从扩展名推断.
        title: 文档标题 (用于 docx/pdf 的元数据).

    Returns:
        实际写入的目标文件路径.
    """
    path = Path(output_path).expanduser().resolve()
    path.parent.mkdir(parents=True, exist_ok=True)

    if fmt is None:
        fmt = infer_format(path)

    # 1. 始终先写 Markdown 文件
    md_path = path.with_suffix(".md")
    md_path.write_text(content, encoding="utf-8")

    # 2. JSON — 直接写目标文件（Markdown 副本已在 md_path）
    if fmt == "json":
        path.write_text(content, encoding="utf-8")
        return path

    # 3. Markdown — 内容已在 md_path，返回实际写入路径
    if fmt == "md":
        return md_path

    # 3. DOCX / PDF — 从 Markdown 文件转换
    success = _convert_via_pandoc(md_path, path, fmt, title=title)
    if success:
        return path

    # 4. 降级
    print(
        "⚠️  Pandoc 不可用，使用内置转换器（输出质量受限）。\n"
        "   建议安装 Pandoc: https://pandoc.org/installing.html",
        file=sys.stderr,
    )
    if fmt == "docx":
        _markdown_to_docx_fallback(content, path, title)
    elif fmt == "pdf":
        _markdown_to_pdf_fallback(content, path, title)

    return path


# ── 降级转换器 (Pandoc 不可用时) ────────────────────────────────


def _markdown_to_docx_fallback(md_text: str, path: Path, title: str = "") -> None:
    """将 Markdown 文本转换为 .docx 文件 (降级转换器).

    仅在 Pandoc 不可用时使用。输出质量受限：不支持粗体、斜体、
    代码块、列表、链接等 Markdown 特性。

    处理规则:
    - # 开头的行 → Heading 样式
    - | 开头的行 → 表格
    - > 开头的行 → 引用 (灰色斜体)
    - 其他 → 正文段落
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise RuntimeError("生成 .docx 需要 python-docx: pip install python-docx")

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # 空行跳过
        if not line.strip():
            i += 1
            continue

        # 标题行
        if line.startswith("#"):
            level = 0
            for ch in line:
                if ch == "#":
                    level += 1
                else:
                    break
            heading_text = line[level:].strip()
            if heading_text:
                doc.add_heading(heading_text, level=min(level, 9))
            i += 1
            continue

        # 表格行 (| 开头)
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            # 跳过分隔行 (如 |---|---|)
            data_rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip().strip("|").split("|")]
                if all(c.replace("-", "").replace(" ", "") == "" for c in cells if c):
                    continue  # 分隔行
                data_rows.append(cells)

            if data_rows:
                cols = max(len(r) for r in data_rows)
                table = doc.add_table(rows=len(data_rows), cols=cols, style="Table Grid")
                for ri, row_cells in enumerate(data_rows):
                    for ci, cell_text in enumerate(row_cells):
                        if ci < cols:
                            table.cell(ri, ci).text = cell_text
                doc.add_paragraph()  # 表后空行
            continue

        # 引用行
        if line.strip().startswith(">"):
            quote_text = line.strip().lstrip("> ").lstrip(">")
            para = doc.add_paragraph()
            run = para.add_run(quote_text)
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            i += 1
            continue

        # 普通段落
        doc.add_paragraph(line)
        i += 1

    # 添加文档标题属性
    if title:
        doc.core_properties.title = title

    doc.save(str(path))


def _markdown_to_pdf_fallback(md_text: str, path: Path, title: str = "") -> None:
    """将 Markdown 文本转换为 .pdf 文件 (降级转换器).

    仅在 Pandoc 不可用时使用。输出质量受限：不支持粗体、斜体、
    代码块、列表等 Markdown 特性，中文支持依赖系统字体。
    """
    try:
        from fpdf import FPDF
    except ImportError:
        raise RuntimeError("生成 .pdf 需要 fpdf2: pip install fpdf2")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # 注册支持中文的字体 (使用系统字体或内置)
    _setup_pdf_font(pdf)

    if title:
        pdf.set_title(title)

    pdf.add_page()

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # 空行
        if not line.strip():
            pdf.ln(4)
            i += 1
            continue

        # 标题行
        if line.startswith("#"):
            level = 0
            for ch in line:
                if ch == "#":
                    level += 1
                else:
                    break
            heading_text = line[level:].strip()
            if heading_text:
                sizes = {1: 18, 2: 15, 3: 13, 4: 12, 5: 11}
                size = sizes.get(level, 10)
                pdf.set_font(style="B", size=size)
                pdf.cell(0, 8, heading_text, new_x="LMARGIN", new_y="NEXT")
                pdf.ln(4)
                pdf.set_font(style="", size=10)
            i += 1
            continue

        # 表格行 — 转为等宽文本
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            pdf.set_font(size=8)
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip().strip("|").split("|")]
                # 跳过纯分隔行
                if all(c.replace("-", "").replace(" ", "") == "" for c in cells if c):
                    continue
                row_text = "  ".join(cells)
                pdf.cell(0, 5, row_text[:120], new_x="LMARGIN", new_y="NEXT")
            pdf.set_font(size=10)
            pdf.ln(2)
            continue

        # 引用行
        if line.strip().startswith(">"):
            quote_text = line.strip().lstrip("> ").lstrip(">")
            pdf.set_font(style="I", size=9)
            pdf.set_text_color(128)
            pdf.multi_cell(0, 5, quote_text)
            pdf.set_text_color(0)
            pdf.set_font(style="", size=10)
            i += 1
            continue

        # 普通段落 — 使用 multi_cell 自动换行
        pdf.set_font(size=10)
        pdf.multi_cell(0, 5, line)
        i += 1

    pdf.output(str(path))


def _setup_pdf_font(pdf) -> None:
    """为 PDF 设置支持中文的字体 (降级转换器用).

    优先尝试系统字体, 降级使用 fpdf2 内置字体。
    """
    import platform

    system = platform.system()

    # macOS 中文字体路径
    if system == "Darwin":
        candidates = [
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/Hiragino Sans GB.ttc",
            "/Library/Fonts/Arial Unicode.ttf",
        ]
    elif system == "Linux":
        candidates = [
            "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/arphic/uming.ttc",
        ]
    else:
        candidates = []

    for font_path in candidates:
        if Path(font_path).exists():
            try:
                pdf.add_font("SystemCN", "", font_path, uni=True)
                pdf.add_font("SystemCN", "B", font_path, uni=True)
                pdf.add_font("SystemCN", "I", font_path, uni=True)
                pdf.set_font("SystemCN", size=10)
                return
            except Exception:
                continue

    # 降级: 使用 fpdf2 内置字体 (不支持中文, 但不会 crash)
    pdf.set_font("Helvetica", size=10)